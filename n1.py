import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog
from PIL import ImageTk, Image

import docx
from docx.oxml.ns import nsdecls
from docxtpl import DocxTemplate
from docx.shared import Pt
import shutil

def Gen_docx_1():
    
    # Create a new Word document
    doc = docx.Document()
    spacing_style = doc.styles.add_style('CustomSpacing', 1)  # Specify the name and base style (1 for 'Normal')
    paragraph_format = spacing_style.paragraph_format
    paragraph_format.space_after = Pt(0)  # Adjust the spacing after the paragraph
    paragraph_format.space_before = Pt(0)  # Adjust the spacing after the paragraph

    # Add a table with three columns and two rows
    table = doc.add_table(rows=5, cols=2)

    hdr_cells = table.rows[0].cells
    hdr_cells[1].text = entry_name.get()
    hdr_cells[1].height = 1

    row_cells = table.rows[1].cells
    row_cells[1].text = 'Email : ', entry_email.get()
    row_cells[1].height = 1

    row_cells = table.rows[2].cells
    row_cells[1].text = 'Course : ', entry_course.get(), ' ', entry_Branch.get()
    row_cells[1].height = 2

    row_cells = table.rows[3].cells
    row_cells[1].text = 'Mobile : ', entry_contact.get()
    row_cells[1].height = 4

    row_cells = table.rows[4].cells
    row_cells[1].text = 'CGPA : ', entry_cgpa.get()

    for row in table.rows:
        for cell in row.cells:
            cell.width = docx.shared.Inches(3)
            cell.height = docx.shared.Inches(2)
            cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph_format = paragraph.paragraph_format
                # paragraph_format.line_spacing = 0.7 # set the line spacing to 0.5
                paragraph_format.space_before = docx.shared.Pt(0) # set the space before to 0pt
                paragraph_format.space_after = docx.shared.Pt(0) # set the space after to 0pt


#Merge the first 5 cells of the first column
    for i in range(4):
        cell = table.cell(i, 0)
        next_cell = table.cell(i + 1, 0)
        cell.merge(next_cell)
# Add a picture to the first cell of the first column
    cell_1_1 = table.cell(0, 0)
    cell_1_1_paragraph = cell_1_1.paragraphs[0]
    cell_1_1_run = cell_1_1_paragraph.add_run()
    cell_1_1_run.add_picture(file_path, width=docx.shared.Inches(1.0))
    cell_1_1.width = docx.shared.Inches(1)

    table1 = doc.add_table(rows=4, cols=5)

# Merge cells in the first row
    cells = table1.rows[0].cells
    cells[0].merge(cells[4])  # Merge the first five cells in the first row
    cells = table1.rows[0].cells
    cells[0].text = 'Academic Details'
    cells[0].height = 1

# Bold the font in the first row
    for cell in table1.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    cells = table1.rows[1].cells
    cells[0].text = 'Course'
    cells[0].height = 1

    cells = table1.rows[1].cells
    cells[1].text = 'Board/University'
    cells[1].height = 1

    cells = table1.rows[1].cells
    cells[2].text = 'Institue'
    cells[2].height = 1

    cells = table1.rows[1].cells
    cells[3].text = 'Year'
    cells[3].height = 1

    cells = table1.rows[1].cells
    cells[4].text = 'Percentage'
    cells[4].height = 1

    for cell in table1.rows[1].cells:
        cell.paragraphs[0].runs[0].bold = True

    cells = table1.rows[2].cells
    cells[0].text = '10th'
    cells[0].height = 2

    cells = table1.rows[2].cells
    cells[1].text = entry_board_10th.get()
    cells[1].height = 2

    cells = table1.rows[2].cells
    cells[2].text = entry_institute_10th.get()
    cells[2].height = 2

    cells = table1.rows[2].cells
    cells[3].text = entry_year_10th.get()
    cells[3].height = 2

    cells = table1.rows[2].cells
    cells[4].text = entry_percentage_10th.get()
    cells[4].height = 2

    cells = table1.rows[3].cells
    cells[0].text = '12th'
    cells[0].height = 2

    cells = table1.rows[3].cells
    cells[1].text = entry_board_12th.get()
    cells[1].height = 2

    cells = table1.rows[3].cells
    cells[2].text = entry_institute_12th.get()
    cells[2].height = 2

    cells = table1.rows[3].cells
    cells[3].text = entry_year_12th.get()
    cells[3].height = 2

    cells = table1.rows[3].cells
    cells[4].text = entry_percentage_12th.get()
    cells[4].height = 2

    # Set the style for the table borders
    table_style = doc.styles['Table Grid']
    table1.style = table_style

# Set the borders for each cell
    for row in table1.rows:
        for cell in row.cells:
            borders = cell._element.xpath('.//w:tcBorders')
            for border in borders:
                for side in border:
                    side.attrib.clear()
                    side.attrib.update(nsdecls('w'), color="000000", space="0", sz="1")

# Add spacing after table1 using the custom style
    doc.add_paragraph(style='CustomSpacing')

    table2 = doc.add_table(rows=2, cols=5)

    # Merge cells in the second, third, and fourth columns for both rows
    for row in table2.rows:
        row.cells[1].merge(row.cells[4])

    # Set the text and height for the cells in the table
    cells = table2.rows[0].cells
    cells[0].text = 'Subects/Electives'
    cells[0].height = 1
    cells[1].text = text_electives.get("1.0", "end-1c")
    cells[1].height = 1

    cells = table2.rows[1].cells
    cells[0].text = 'Technical Proficiency'
    cells[0].height = 1
    cells[1].text = text_technical_proficiency.get("1.0", "end-1c")
    cells[1].height = 1

    for cell in table2.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        break
    for cell in table2.rows[1].cells:
        cell.paragraphs[0].runs[0].bold = True
        break



# Set the style for the table borders
    table_style = doc.styles['Table Grid']
    table2.style = table_style

# Set the borders for each cell
    for row in table2.rows:
        for cell in row.cells:
            borders = cell._element.xpath('.//w:tcBorders')
            for border in borders:
                for side in border:
                    side.attrib.clear()
                    side.attrib.update(nsdecls('w'), color="000000", space="0", sz="1")
    # Add spacing after table1 using the custom style
    doc.add_paragraph(style='CustomSpacing')

    table3 = doc.add_table(rows=2, cols=1)
    # Set the content and formatting of the cells
    cells = table3.rows[0].cells
    cells[0].text = 'Summer Internship/Work Experience'
    cells[0].height = 1

    cells = table3.rows[1].cells
    cells[0].text = text_internship.get("1.0", "end-1c")
    cells[0].height = 1
    for cell in table3.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    table_style = doc.styles['Table Grid']
    table3.style = table_style

# Set the borders for each cell
    for row in table3.rows:
        for cell in row.cells:
            borders = cell._element.xpath('.//w:tcBorders')
            for border in borders:
                for side in border:
                    side.attrib.clear()
                    side.attrib.update(nsdecls('w'), color="000000", space="0", sz="1")

# Add spacing after table1 using the custom style
    doc.add_paragraph(style='CustomSpacing')

    table9 = doc.add_table(rows=2, cols=1)
# Set the content and formatting of the cells
    cells = table9.rows[0].cells
    cells[0].text = 'Projects'
    cells[0].height = 1

    cells = table9.rows[1].cells
    cells[0].text = text_project_1.get("1.0", "end-1c") ,"\n", text_project_2.get("1.0", "end-1c")
    cells[0].height = 1
    for cell in table9.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True



# Set the style for the table borders
    table_style = doc.styles['Table Grid']
    table9.style = table_style

# Set the borders for each cell
    for row in table9.rows:
        for cell in row.cells:
            borders = cell._element.xpath('.//w:tcBorders')
            for border in borders:
                for side in border:
                    side.attrib.clear()
                    side.attrib.update(nsdecls('w'), color="000000", space="0", sz="1")

# Add spacing after table1 using the custom style
    doc.add_paragraph(style='CustomSpacing')

# Set the style for the table borders


    table4 = doc.add_table(rows=2, cols=1)
# Set the content and formatting of the cells
    cells = table4.rows[0].cells
    cells[0].text = 'Position of responsibilities'
    cells[0].height = 1

    cells = table4.rows[1].cells
    cells[0].text = text_position_responsibility.get("1.0", "end-1c")
    cells[0].height = 1
    for cell in table4.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True



# Set the style for the table borders
    table_style = doc.styles['Table Grid']
    table4.style = table_style

# Set the borders for each cell
    for row in table4.rows:
        for cell in row.cells:
            borders = cell._element.xpath('.//w:tcBorders')
            for border in borders:
                for side in border:
                    side.attrib.clear()
                    side.attrib.update(nsdecls('w'), color="000000", space="0", sz="1")
# Add spacing after table1 using the custom style
    doc.add_paragraph(style='CustomSpacing')

    table5 = doc.add_table(rows=2, cols=1)
# Set the content and formatting of the cells
    cells = table5.rows[0].cells
    cells[0].text = 'Interrest'
    cells[0].height = 1

    cells = table5.rows[1].cells
    cells[0].text = text_interest.get("1.0", "end-1c")
    cells[0].height = 1
    for cell in table5.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    # Set the style for the table borders
    table_style = doc.styles['Table Grid']
    table5.style = table_style

    # Set the borders for each cell
    for row in table5.rows:
        for cell in row.cells:
            borders = cell._element.xpath('.//w:tcBorders')
            for border in borders:
                for side in border:
                    side.attrib.clear()
                    side.attrib.update(nsdecls('w'), color="000000", space="0", sz="1")
# Add spacing after table1 using the custom style
    doc.add_paragraph(style='CustomSpacing')

    table6 = doc.add_table(rows=2, cols=1)
# Set the content and formatting of the cells
    cells = table6.rows[0].cells
    cells[0].text = 'Awards and Recognition'
    cells[0].height = 1

    cells = table6.rows[1].cells
    cells[0].text = text_awards_recognition.get("1.0", "end-1c")
    cells[0].height = 1
    for cell in table6.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True



    # Set the style for the table borders
    table_style = doc.styles['Table Grid']
    table6.style = table_style

# Set the borders for each cell
    for row in table6.rows:
        for cell in row.cells:
            borders = cell._element.xpath('.//w:tcBorders')
            for border in borders:
                for side in border:
                    side.attrib.clear()
                    side.attrib.update(nsdecls('w'), color="000000", space="0", sz="1")
# Add spacing after table1 using the custom style
    doc.add_paragraph(style='CustomSpacing')

    table7 = doc.add_table(rows=2, cols=1)
# Set the content and formatting of the cells
    cells = table7.rows[0].cells
    cells[0].text = 'Volunteer Experience'
    cells[0].height = 1

    cells = table7.rows[1].cells
    cells[0].text = text_volunteer_experience.get("1.0", "end-1c")
    cells[0].height = 1
    for cell in table7.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True

    # Set the style for the table borders
    table_style = doc.styles['Table Grid']
    table7.style = table_style

# Set the borders for each cell
    for row in table7.rows:
        for cell in row.cells:
            borders = cell._element.xpath('.//w:tcBorders')
            for border in borders:
                for side in border:
                    side.attrib.clear()
                    side.attrib.update(nsdecls('w'), color="000000", space="0", sz="1")
# Add spacing after table1 using the custom style
    doc.add_paragraph(style='CustomSpacing')

    table8 = doc.add_table(rows=2, cols=1)
# Set the content and formatting of the cells
    cells = table8.rows[0].cells
    cells[0].text = 'Language Known'
    cells[0].height = 1

    cells = table8.rows[1].cells
    cells[0].text = text_languages_known.get("1.0", "end-1c")
    cells[0].height = 1

    for cell in table8.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True


# Set the style for the table borders
    table_style = doc.styles['Table Grid']
    table8.style = table_style

# Set the borders for each cell
    for row in table8.rows:
        for cell in row.cells:
            borders = cell._element.xpath('.//w:tcBorders')
            for border in borders:
                for side in border:
                    side.attrib.clear()
                    side.attrib.update(nsdecls('w'), color="000000", space="0", sz="1")

# Save the document
    doc.save('Resume.docx')

def download_file_1():
    Gen_docx_1()

    # Get the source file path in the current working directory
    source_file_path = "./Resume.docx"

    # Ask the user to choose the download location
    file_path1 = filedialog.asksaveasfilename(defaultextension=".docx", initialfile="Resume.docx")

    # Append the extension if not already present
    if not file_path1.endswith(".docx"):
        file_path1 += ".docx"

    # Copy the file to the chosen location
    shutil.copy(source_file_path, file_path1)

def Gen_docx_2():
    doc = DocxTemplate("Resume-1.docx")
    data = {"NAME": entry_name.get()}

    doc.render(data)
    doc.save("Resume_template.docx")

def download_file_2():
    # Get the source file path in the current working directory
    Gen_docx_1()

    source_file_path2 = "./Resume_template.docx"

    # Ask the user to choose the download location
    file_path2 = filedialog.asksaveasfilename(defaultextension=".docx", initialfile="Resume.docx")

    # Append the extension if not already present
    if not file_path2.endswith(".docx"):
        file_path2 += ".docx"

    # Copy the file to the chosen location
    shutil.copy(source_file_path2, file_path2)

def save_data():
    name = entry_name.get()
    email = entry_email.get()
    contact = entry_contact.get()
    course = entry_course.get()
    cgpa = entry_cgpa.get()
    college = entry_clg.get()
    branch = entry_Branch.get()
    interest = text_interest.get("1.0", "end-1c")
    board_10th = entry_board_10th.get()
    year_10th = entry_year_10th.get()
    institute_10th = entry_institute_10th.get()
    percentage_10th = entry_percentage_10th.get()
    board_12th = entry_board_12th.get()
    year_12th = entry_year_12th.get()
    institute_12th = entry_institute_12th.get()
    percentage_12th = entry_percentage_12th.get()
    electives = text_electives.get("1.0", "end-1c")
    technical_proficiency = text_technical_proficiency.get("1.0", "end-1c")
    internship = text_internship.get("1.0", "end-1c")
    project_1 = text_project_1.get("1.0", "end-1c")
    project_2 = text_project_2.get("1.0", "end-1c")
    project_3 = text_project_3.get("1.0", "end-1c")
    project_4 = text_project_4.get("1.0", "end-1c")
    position_responsibility = text_position_responsibility.get("1.0", "end-1c")
    awards_recognition = text_awards_recognition.get("1.0", "end-1c")
    volunteer_experience = text_volunteer_experience.get("1.0", "end-1c")
    languages_known = text_languages_known.get("1.0", "end-1c")


    resume_data.append({
        "Name": name,
        "Email": email,
        "Contact": contact,
        "Course": course,
        "CGPA": cgpa,
        "College": college,
        "Branch": branch,
        # "Address": address,
        "Interest": interest,
        "10th Board": board_10th,
        "Year of Passing (10th)": year_10th,
        "Institute (10th)": institute_10th,
        "Percentage (10th)": percentage_10th,
        "12th Board": board_12th,
        "Year of Passing (12th)": year_12th,
        "Institute (12th)": institute_12th,
        "Percentage (12th)": percentage_12th,
        "Electives": electives,
        "Technical Proficiency": technical_proficiency,
        "Internship": internship,
        "Project 1": project_1,
        "Project 2": project_2,
        "Project 3": project_3,
        "Project 4": project_4,
        "Position of Responsibility": position_responsibility,
        "Awards and Recognition": awards_recognition,
        "Volunteer Experience": volunteer_experience,
        "Languages Known": languages_known
        # "Hobbies": hobbies,
        # "Linkedin": linkedin

    })

    messagebox.showinfo("Information", "Resume data saved successfully!")

root = tk.Tk()
root.title("Resume Form")
root.state('zoomed')

bg=ImageTk.PhotoImage(file = 'bk.png')
bg_image = tk.Label(root , image=bg)
bg_image.pack()


label = tk.Label(root, text="Personal Information", width=18, font=("bold", 18))
label.place(x=350,y=10)
resume_data = []

default_image = Image.open("default_image.png")
resized_default_image = default_image.resize((150,150))
default_photo = ImageTk.PhotoImage(resized_default_image)

image_label = tk.Label(root, image=default_photo)
image_label.image = default_photo
image_label.place(x=500, y=70)

def choose_image():
       global file_path
       file_path = filedialog.askopenfilename(filetypes=[("Image files", "*.jpg;*.jpeg;*.png")])
       if file_path:
            image = Image.open(file_path)
            resized_image = image.resize((150,150))
            photo = ImageTk.PhotoImage(resized_image)
            image_label.configure(image=photo)
            image_label.image = photo
       else:
           image = Image.open("default_profile.png")
           resized_image = default_image.resize((150,150))
           photo = ImageTk.PhotoImage(resized_image)
           image_label.configure(image=photo)
           image_label.image = photo
           
choose_image_btn = tk.Button(root, text="Choose Image", command=choose_image)
choose_image_btn.place(x=535, y=250)

# Label and Entry fields
label_name = tk.Label(root, text="Full Name:", width=18, font=("bold", 11))
label_name.place(x=80,y=50)#grid(row=0, column=0, padx=10, pady=5, sticky="e")
entry_name = tk.Entry(root)
entry_name.place(x=260,y=50)

#  Add labels and entry fields for email, contact, address, interest, course, cgpa, 10th board, year of passing, institute, percentage
label_email = tk.Label(root, text="Email:", width=18, font=("bold", 11))
label_email.place(x=80,y=80)#grid(row=1, column=0, padx=10, pady=5, sticky="e")
entry_email = tk.Entry(root)
entry_email.place(x=260,y=80)#grid(row=1, column=1, padx=10, pady=5, sticky="w")

label_contact = tk.Label(root, text="Contact:", width=18, font=("bold", 11))
label_contact.place(x=80,y=110)#grid(row=1, column=2, padx=10, pady=5, sticky="e")
entry_contact = tk.Entry(root)
entry_contact.place(x=260,y=110)#grid(row=1, column=3, padx=10, pady=5, sticky="w")

label_course = tk.Label(root, text="Course:", width=18, font=("bold", 11))
label_course.place(x=80,y=150)#grid(row=2, column=0, padx=10, pady=5, sticky="e")
entry_course = tk.Entry(root)
entry_course.place(x=260,y=150)#grid(row=2, column=1, padx=10, pady=5, sticky="w")

label_cgpa = tk.Label(root, text="CGPA:", width=18, font=("bold", 11))
label_cgpa.place(x=80,y=180)#grid(row=2, column=2, padx=10, pady=5, sticky="e")
entry_cgpa = tk.Entry(root)
entry_cgpa.place(x=260,y=180)#grid(row=2, column=3, padx=10, pady=5, sticky="w")

label_clg = tk.Label(root, text="Insitue Name:", width=18, font=("bold", 11))
label_clg.place(x=80,y=210)#grid(row=3, column=0, padx=10, pady=5, sticky="e")
entry_clg = tk.Entry(root)
entry_clg.place(x=260,y=210)#grid(row=3, column=1, padx=10, pady=5, sticky="w")

label_Branch = tk.Label(root, text="Branch:", width=18, font=("bold", 11))
label_Branch.place(x=80,y=240)#grid(row=3, column=2, padx=10, pady=5, sticky="e")
entry_Branch = tk.Entry(root)
entry_Branch.place(x=260,y=240)#grid(row=3, column=3, padx=10, pady=5, sticky="w")

label_interest = tk.Label(root, text="Interest:", width=18, font=("bold", 11))
label_interest.place(x=80,y=280)#grid(row=4, column=2, padx=10, pady=5, sticky="e")
text_interest = tk.Text(root, height=3, width=30)
text_interest.place(x=260,y=280)#grid(row=4, column=3, padx=10, pady=5, sticky="w")

label = tk.Label(root,text="Academic Information", width=20, font=("bold", 18))
label.place(x=1000,y=10)

label_board_10th = tk.Label(root, text="10th Board:", width=18, font=("bold", 11))
label_board_10th.place(x=800,y=70)#grid(row=5, column=0, padx=10, pady=5, sticky="e")
entry_board_10th = tk.Entry(root)
entry_board_10th.place(x=980,y=70,height=25)#grid(row=5, column=1, padx=10, pady=5, sticky="w")

label_year_10th = tk.Label(root, text="Year of Passing (10th):", width=18, font=("bold", 11))
label_year_10th.place(x=800,y=110)#grid(row=6, column=0, padx=10, pady=5, sticky="e")
entry_year_10th = tk.Entry(root)
entry_year_10th.place(x=980,y=110,height=25)#grid(row=6, column=1, padx=10, pady=5, sticky="w")

label_institute_10th = tk.Label(root, text="Institute (10th):", width=18, font=("bold", 11))
label_institute_10th.place(x=800,y=150)#grid(row=7, column=0, padx=10, pady=5, sticky="e")
entry_institute_10th = tk.Entry(root)
entry_institute_10th.place(x=980,y=150,height=25)#grid(row=7, column=1, padx=10, pady=5, sticky="w")

label_percentage_10th = tk.Label(root, text="Percentage (10th):", width=18, font=("bold", 11))
label_percentage_10th.place(x=800,y=190)#grid(row=8, column=0, padx=10, pady=5, sticky="e")
entry_percentage_10th = tk.Entry(root)
entry_percentage_10th.place(x=980,y=190,height=25)#grid(row=8, column=1, padx=10, pady=5, sticky="w")

label_board_12th = tk.Label(root, text="12th Board:", width=18, font=("bold", 11))
label_board_12th.place(x=1200,y=70)#grid(row=5, column=2, padx=10, pady=5, sticky="e")
entry_board_12th = tk.Entry(root)
entry_board_12th.place(x=1380,y=70,height=25)#grid(row=5, column=3, padx=10, pady=5, sticky="w")

label_year_12th = tk.Label(root, text="Year of Passing (12th):", width=18, font=("bold", 11))
label_year_12th.place(x=1200,y=110)#grid(row=6, column=2, padx=10, pady=5, sticky="e")
entry_year_12th = tk.Entry(root)
entry_year_12th.place(x=1380,y=110,height=25)#grid(row=6, column=3, padx=10, pady=5, sticky="w")

label_institute_12th = tk.Label(root, text="Institute (12th):", width=18, font=("bold", 11))
label_institute_12th.place(x=1200,y=150)#grid(row=7, column=2, padx=10, pady=5, sticky="e")
entry_institute_12th = tk.Entry(root)
entry_institute_12th.place(x=1380,y=150,height=25)#grid(row=7, column=3, padx=10, pady=5, sticky="w")

label_percentage_12th = tk.Label(root, text="Percentage (12th):", width=18, font=("bold", 11))
label_percentage_12th.place(x=1200,y=190)#grid(row=8, column=2, padx=10, pady=5, sticky="e")
entry_percentage_12th = tk.Entry(root)
entry_percentage_12th.place(x=1380,y=190,height=25)#grid(row=8, column=3, padx=10, pady=5, sticky="w")

label = tk.Label(root,text="Co-curricular Information", width=20, font=("bold", 18))
label.place(x=650,y=335)

# Additional fields
label_electives = tk.Label(root, text="Subjects/Electives:", width=18, font=("bold", 11))
label_electives.place(x=80,y=400)#grid(row=9, column=0, padx=10, pady=5, sticky="e")
text_electives = tk.Text(root, height=2, width=30)
text_electives.place(x=260,y=400)#grid(row=9, column=1, padx=10, pady=5, sticky="w")

label_technical_proficiency = tk.Label(root, text="Technical Proficiency:", width=18, font=("bold", 11))
label_technical_proficiency.place(x=80,y=450)#grid(row=9, column=2, padx=10, pady=5, sticky="e")
text_technical_proficiency = tk.Text(root, height=2, width=30)
text_technical_proficiency.place(x=260,y=450)#grid(row=9, column=3, padx=10, pady=5, sticky="w")

label_internship = tk.Label(root, text="Internship:", width=18, font=("bold", 11))
label_internship.place(x=80,y=500)#grid(row=10, column=0, padx=10, pady=5, sticky="e")
text_internship = tk.Text(root, height=3, width=30)
text_internship.place(x=260,y=500)#grid(row=10, column=1, padx=10, pady=5, sticky="w")

label_experience = tk.Label(root, text="Experience", width=18, font=("bold", 11))
label_experience.place(x=80,y=550)#grid(row=10, column=2, padx=10, pady=5, sticky="e")
text_experience = tk.Text(root, height=3, width=30)
text_experience.place(x=260,y=550)#grid(row=10, column=3, padx=10, pady=5, sticky="w")

label_project_1 = tk.Label(root, text="Project 1:", width=18, font=("bold", 11))
label_project_1.place(x=518,y=400)#grid(row=11, column=0, padx=10, pady=5, sticky="e")
text_project_1 = tk.Text(root, height=3, width=30)
text_project_1.place(x=680,y=400)#grid(row=11, column=1, padx=10, pady=5, sticky="w")

label_project_2 = tk.Label(root, text="Project 2:", width=18, font=("bold", 11))
label_project_2.place(x=520,y=450)#grid(row=11, column=2, padx=10, pady=5, sticky="e")
text_project_2 = tk.Text(root, height=3, width=30)
text_project_2.place(x=680,y=450)#grid(row=11, column=3, padx=10, pady=5, sticky="w")

label_project_3 = tk.Label(root, text="Project 3:", width=18, font=("bold", 11))
label_project_3.place(x=520,y=500)#grid(row=12, column=0, padx=10, pady=5, sticky="e")
text_project_3 = tk.Text(root, height=3, width=30)
text_project_3.place(x=680,y=500)#grid(row=12, column=1, padx=10, pady=5, sticky="w")

label_project_4 = tk.Label(root, text="Project 4:", width=18, font=("bold", 11))
label_project_4.place(x=520,y=550)#grid(row=12, column=2, padx=10, pady=5, sticky="e")
text_project_4 = tk.Text(root, height=3, width=30)
text_project_4.place(x=680,y=550)#grid(row=12, column=3, padx=10, pady=5, sticky="w")

label_position_responsibility = tk.Label(root, text="Position of Responsibility:", width=18, font=("bold", 11))
label_position_responsibility.place(x=1000,y=400)#grid(row=13, column=0, padx=10, pady=5, sticky="e")
text_position_responsibility = tk.Text(root, height=2, width=30)
text_position_responsibility.place(x=1180,y=400)#grid(row=13, column=1, padx=10, pady=5, sticky="w")

label_awards_recognition = tk.Label(root, text="Awards and Recognition:", width=18, font=("bold", 11))
label_awards_recognition.place(x=1000,y=450)#grid(row=13, column=2, padx=10, pady=5, sticky="e")
text_awards_recognition = tk.Text(root, height=2, width=30)
text_awards_recognition.place(x=1180,y=450)#grid(row=13, column=3, padx=10, pady=5, sticky="w")

label_volunteer_experience = tk.Label(root, text="Volunteer Experience:", width=18, font=("bold", 11))
label_volunteer_experience.place(x=1000,y=500)#grid(row=14, column=0, padx=10, pady=5, sticky="e")
text_volunteer_experience = tk.Text(root, height=2, width=30)
text_volunteer_experience.place(x=1180,y=500)#grid(row=14, column=1, padx=10, pady=5, sticky="w")

label_languages_known = tk.Label(root, text="Languages Known:", width=18, font=("bold", 11))
label_languages_known.place(x=1000,y=550)#grid(row=14, column=2, padx=10, pady=5, sticky="e")
text_languages_known = tk.Text(root, height=2, width=30)
text_languages_known.place(x=1180,y=550)#grid(row=14, column=3, padx=10, pady=5, sticky="w")

# Submit and Download buttons
button_save = tk.Button(root, text="Submit", width=10, font=("bold", 11), bg="#CCFFFF", command=save_data)
button_save.place(x=700,y=620)

download1 = tk.Label(root,text="Superset Format", width=15, font=("bold", 11))
download1.place(x=580,y=685)

button_download1 = tk.Button(root, text="Download", width=10, font=("bold", 11),bg="#E0FFFF", command=download_file_1)
button_download1.place(x=750,y=680)

root.mainloop()